from docx import Document
from docx.shared import Pt

class Docx:
    
    doc = Document()

    def __init__(self):
        super().__init__()
    
    def exercise_1(self):

        self.doc.add_heading('This heading')
        self.doc.add_paragraph('lorem text')
        self.doc.add_paragraph('lorem text again')

        self.doc.save('Ex.docx')

    def exercise_2(self):

        self.doc.add_paragraph('PARAGRAPH 1')
        self.doc.add_page_break()
        self.doc.add_paragraph('PARAGRAPH 2')
        self.doc.add_page_break()
        self.doc.add_paragraph('PARAGRAPH 3')

        self.doc.save('Ex.docx')

    def exercise_3(self):

        document = Document('Ex.docx')

        document.add_paragraph().add_run('text not style')
        document.add_paragraph().add_run('bold text').bold = True
        document.add_paragraph().add_run('text not style again')
        document.add_paragraph().add_run('bold text').italic = True

        for i in document.paragraphs:
            print(i.text)

        document.save('Ex.docx')

    def exercise_4(self):

        document = Document()
        document.add_paragraph().add_run('test for working!').add_picture('sample.png')

        document.save('Ex.docx')


    def exercise_5(self):

        text_font = self.doc.add_paragraph('first text ').add_run('test text for font!')
        font = text_font.font
        font.name = 'system'
        font.size = Pt(22)

        text_font = self.doc.add_paragraph('first text ').add_run('test text for font!')
        font = text_font.font
        font.name = 'Roboto'
        font.size = Pt(22)

        text_font = self.doc.add_paragraph('first text ').add_run('test text for font!')
        font = text_font.font
        font.name = 'Britannin'
        font.size = Pt(22)

        self.doc.save('Ex.docx')

if __name__ == "__main__":
    main = Docx()
    main.exercise_2()